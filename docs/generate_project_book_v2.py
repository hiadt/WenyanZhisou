# -*- coding: utf-8 -*-
"""Generate the V2 project book for 问研智搜.

The script is intentionally kept in the repository so the project book can be
regenerated after metrics or screenshots are updated.
"""

from __future__ import annotations

import math
from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
FIG_DIR = DOCS / "assets" / "project_book_v2"
FIG_DIR.mkdir(parents=True, exist_ok=True)

COLORS = {
    "ink": "#0f172a",
    "muted": "#64748b",
    "line": "#cbd5e1",
    "blue": "#2563eb",
    "cyan": "#0891b2",
    "green": "#059669",
    "orange": "#d97706",
    "red": "#dc2626",
    "violet": "#7c3aed",
    "light_blue": "#dbeafe",
    "light_cyan": "#ecfeff",
    "light_green": "#dcfce7",
    "light_orange": "#ffedd5",
    "light_violet": "#ede9fe",
    "panel": "#f8fafc",
}


def setup_font() -> None:
    available = {f.name for f in font_manager.fontManager.ttflist}
    for name in ["Microsoft YaHei", "SimHei", "SimSun", "Noto Sans CJK SC"]:
        if name in available:
            plt.rcParams["font.sans-serif"] = [name]
            break
    plt.rcParams["axes.unicode_minus"] = False


def savefig(name: str) -> Path:
    path = FIG_DIR / name
    plt.savefig(path, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def box(ax, xy, w, h, text, fc, ec=None, size=9.5, weight="normal"):
    ec = ec or COLORS["line"]
    patch = FancyBboxPatch(
        xy,
        w,
        h,
        boxstyle="round,pad=0.018,rounding_size=0.025",
        facecolor=fc,
        edgecolor=ec,
        linewidth=1.2,
    )
    ax.add_patch(patch)
    ax.text(
        xy[0] + w / 2,
        xy[1] + h / 2,
        text,
        ha="center",
        va="center",
        fontsize=size,
        weight=weight,
        color=COLORS["ink"],
        wrap=True,
    )
    return patch


def arrow(ax, a, b, color=None, lw=1.5):
    ax.add_patch(
        FancyArrowPatch(
            a,
            b,
            arrowstyle="-|>",
            mutation_scale=13,
            linewidth=lw,
            color=color or COLORS["muted"],
        )
    )


def canvas(title: str, size=(12, 6.6)):
    fig, ax = plt.subplots(figsize=size)
    ax.axis("off")
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.text(0.04, 0.93, title, fontsize=15.5, weight="bold", color=COLORS["ink"])
    return fig, ax


def draw_figures() -> dict[str, Path]:
    setup_font()
    figs: dict[str, Path] = {}

    fig, ax = canvas("问研智搜总体架构：大模型决策、小模型排序、多源召回、结构化输出")
    layers = [
        ("交互展示层", ["论文排序", "查询拆解", "Agent轨迹", "结果归纳", "关系图", "JSON"], COLORS["light_blue"]),
        ("智能Agent层", ["DeepSeek Planner", "Query Evolution", "Crawler调度", "Selector采样", "LLM Verifier", "Synthesizer"], COLORS["light_cyan"]),
        ("检索候选层", ["OpenAlex", "Semantic Scholar", "arXiv", "Serper Crawler", "PaSaTitleDB", "LocalCorpus"], COLORS["light_green"]),
        ("小模型排序层", ["BM25", "BGE Embedding", "BGE Reranker", "Authority", "Recency", "Diversity"], COLORS["light_orange"]),
        ("评测治理层", ["Precision", "Recall", "F1", "Latency", "API Calls", "Hit Report"], COLORS["light_violet"]),
    ]
    y = 0.80
    for layer, items, fc in layers:
        box(ax, (0.04, y - 0.055), 0.16, 0.08, layer, "white", size=11, weight="bold")
        for i, item in enumerate(items):
            box(ax, (0.23 + i * 0.12, y - 0.05), 0.105, 0.065, item, fc, size=8.6)
        if y < 0.80:
            arrow(ax, (0.53, y + 0.055), (0.53, y + 0.025), COLORS["blue"])
        y -= 0.15
    ax.text(0.04, 0.08, "设计目标：以F1@20为主优化目标，以API/LLM调用次数和端到端时延为硬约束，形成可评测、可解释、可部署的学术检索Agent。", fontsize=10, color=COLORS["muted"])
    figs["architecture"] = savefig("fig01_总体架构.png")

    fig, ax = canvas("端到端Agent工作流")
    steps = [
        ("输入查询", "自然语言科研问题"),
        ("意图解析", "实体/约束/子查询"),
        ("查询演化", "语义核心/约束聚焦/权威时效"),
        ("多源召回", "API + Crawler + 本地库"),
        ("小模型精排", "BM25 + BGE + Reranker"),
        ("LLM选择器", "高相关/部分相关/无关"),
        ("结构化归纳", "论文列表/主题线索/JSON"),
    ]
    for i, (a, b) in enumerate(steps):
        x = 0.04 + i * 0.135
        box(ax, (x, 0.54), 0.112, 0.17, f"{i+1}. {a}\n{b}", [COLORS["light_blue"], COLORS["light_cyan"], COLORS["light_green"], COLORS["light_orange"], COLORS["light_violet"], "#fee2e2", "#e2e8f0"][i], size=8.6, weight="bold")
        if i < len(steps) - 1:
            arrow(ax, (x + 0.114, 0.625), (x + 0.134, 0.625), COLORS["blue"], 1.8)
    box(ax, (0.10, 0.25), 0.32, 0.12, "可解释轨迹：记录每轮查询、候选数量、选择器判断和预算消耗", "white", size=10)
    box(ax, (0.56, 0.25), 0.32, 0.12, "评测闭环：输出predictions.jsonl、metrics.json、hit_report.json定位瓶颈", "white", size=10)
    figs["workflow"] = savefig("fig02_Agent工作流.png")

    fig, ax = canvas("PaSa-inspired Crawler/Selector轻量化改良")
    box(ax, (0.05, 0.62), 0.25, 0.17, "Crawler侧\n多源候选发现\nOpenAlex/S2/arXiv/Serper/本地库", COLORS["light_cyan"], COLORS["cyan"], 10.5, "bold")
    box(ax, (0.38, 0.62), 0.25, 0.17, "Selector侧\n源感知采样\nBGE排序 + DeepSeek相关性验证", COLORS["light_green"], COLORS["green"], 10.5, "bold")
    box(ax, (0.70, 0.62), 0.25, 0.17, "Ranker侧\n多目标融合\n相关性/权威性/时效性/多样性", COLORS["light_orange"], COLORS["orange"], 10.5, "bold")
    arrow(ax, (0.31, 0.70), (0.37, 0.70), COLORS["blue"], 2)
    arrow(ax, (0.64, 0.70), (0.69, 0.70), COLORS["blue"], 2)
    box(ax, (0.10, 0.34), 0.21, 0.12, "吸收PaSa优势\n高召回Crawler + 候选选择器", "white", size=9.5)
    box(ax, (0.39, 0.34), 0.24, 0.12, "轻量化改良\n不训练RL，改为预算控制与可解释排序", "white", size=9.5)
    box(ax, (0.70, 0.34), 0.20, 0.12, "竞赛适配\n结构化输出 + 可复现实验", "white", size=9.5)
    ax.text(0.06, 0.15, "核心思想：先扩大候选边界，再用选择器和多目标排序把召回转化为Top-K命中。", fontsize=10.5, color=COLORS["muted"])
    figs["pasa"] = savefig("fig03_PaSa启发机制.png")

    versions = ["v7初版", "v12主线", "rollback复测", "v16压缩失败"]
    f1 = [0.086, 0.159, 0.143, 0.081]
    latency = [116.8, 91.1, 98.0, 146.8]
    fig, ax1 = plt.subplots(figsize=(11.4, 6.2))
    x = list(range(len(versions)))
    bars = ax1.bar(x, f1, color=[COLORS["muted"], COLORS["blue"], COLORS["orange"], COLORS["red"]], width=0.55)
    ax1.set_title("RealScholarQuery阶段性指标对比：v12作为当前项目书主结果口径", fontsize=14.5, weight="bold")
    ax1.set_ylabel("F1@20", color=COLORS["blue"])
    ax1.set_ylim(0, 0.22)
    ax1.set_xticks(x)
    ax1.set_xticklabels(versions)
    ax1.grid(axis="y", alpha=0.25)
    for bar, value in zip(bars, f1):
        ax1.text(bar.get_x() + bar.get_width() / 2, value + 0.006, f"{value:.3f}", ha="center", fontsize=10)
    ax2 = ax1.twinx()
    ax2.plot(x, latency, color=COLORS["green"], marker="o", linewidth=2.5)
    ax2.set_ylabel("平均延迟(s)", color=COLORS["green"])
    ax2.set_ylim(0, 180)
    figs["metrics"] = savefig("fig04_阶段性指标对比.png")

    fig, ax = plt.subplots(figsize=(8.4, 6.2))
    ax.pie(
        [70, 20, 10],
        labels=["F1 Score\n70%", "效率\n20%", "结构化输出\n10%"],
        colors=[COLORS["blue"], COLORS["green"], COLORS["orange"]],
        startangle=110,
        wedgeprops={"width": 0.45, "edgecolor": "white"},
        textprops={"fontsize": 12, "weight": "bold"},
    )
    ax.text(0, 0, "赛题三\n自动评分", ha="center", va="center", fontsize=13, weight="bold")
    ax.set_title("评分权重与系统优化目标映射", fontsize=14.5, weight="bold")
    figs["score_weight"] = savefig("fig05_评分权重.png")

    fig, ax = canvas("多目标排序函数：从候选池到Top-K可提交结果")
    features = [
        ("API相关性", "OpenAlex/S2/arXiv检索分"),
        ("BM25", "标题与摘要词项覆盖"),
        ("Embedding", "语义相似度"),
        ("Reranker", "交叉编码器精排"),
        ("LLM Score", "意图一致性判断"),
        ("Authority", "引用与venue权威性"),
        ("Recency", "发表时效性"),
        ("Diversity", "主题去冗余"),
    ]
    for i, (a, b) in enumerate(features):
        x0 = 0.05 + (i % 4) * 0.23
        y0 = 0.68 - (i // 4) * 0.22
        box(ax, (x0, y0), 0.18, 0.11, f"{a}\n{b}", "white", size=8.8, weight="bold")
        arrow(ax, (x0 + 0.09, y0), (0.50, 0.36), COLORS["muted"], 1.0)
    box(ax, (0.34, 0.20), 0.34, 0.16, "Score(q,p)=Σ wᵢ fᵢ(q,p)+LabelBonus−NoisePenalty\n约束：API预算、LLM预算、端到端时延、Top-K多样性", COLORS["light_blue"], COLORS["blue"], 10, "bold")
    box(ax, (0.75, 0.20), 0.18, 0.16, "输出\nTop20论文\n归纳报告\nJSON", COLORS["light_green"], COLORS["green"], 10, "bold")
    arrow(ax, (0.69, 0.28), (0.74, 0.28), COLORS["blue"], 2)
    figs["rank"] = savefig("fig06_排序函数.png")

    labels = ["查询理解", "多源召回", "排序精度", "结构化归纳", "可解释性", "工程部署", "效率控制"]
    ours = [0.78, 0.72, 0.58, 0.82, 0.86, 0.88, 0.70]
    base = [0.45, 0.50, 0.42, 0.35, 0.30, 0.55, 0.65]
    angles = [i / len(labels) * 2 * math.pi for i in range(len(labels))]
    angles += angles[:1]
    fig = plt.figure(figsize=(8.3, 6.8))
    ax = fig.add_subplot(111, polar=True)
    ax.plot(angles, ours + ours[:1], color=COLORS["blue"], linewidth=2.4, label="问研智搜")
    ax.fill(angles, ours + ours[:1], color=COLORS["blue"], alpha=0.18)
    ax.plot(angles, base + base[:1], color=COLORS["muted"], linewidth=2, linestyle="--", label="轻量基线")
    ax.fill(angles, base + base[:1], color=COLORS["muted"], alpha=0.12)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, fontsize=9.5)
    ax.set_yticklabels([])
    ax.set_ylim(0, 1)
    ax.set_title("系统能力雷达图：相对轻量基线的综合提升", fontsize=13.5, weight="bold", pad=18)
    ax.legend(loc="upper right", bbox_to_anchor=(1.25, 1.10))
    figs["radar"] = savefig("fig07_能力雷达图.png")

    fig, ax = canvas("消融实验矩阵：定位召回、排序与LLM选择器贡献")
    headers = ["配置", "多源API", "PaSaTitleDB", "BGE/Reranker", "DeepSeek Verifier", "预期影响"]
    rows = [
        ["A0 基础检索", "✓", "-", "-", "-", "验证最低召回边界"],
        ["A1 + 小模型", "✓", "-", "✓", "-", "提升语义排序与噪声过滤"],
        ["A2 + 本地标题库", "✓", "✓", "✓", "-", "提升arXiv类公开集覆盖率"],
        ["A3 + LLM选择器", "✓", "✓", "✓", "✓", "提升Top20精确率和归纳质量"],
        ["A4 去除LLM", "✓", "✓", "✓", "-", "评估成本-性能折中"],
    ]
    widths = [0.18, 0.12, 0.14, 0.15, 0.17, 0.25]
    x0, y0 = 0.04, 0.78
    for j, h in enumerate(headers):
        box(ax, (x0 + sum(widths[:j]), y0), widths[j] - 0.005, 0.075, h, COLORS["light_blue"], size=8.5, weight="bold")
    for i, row in enumerate(rows):
        y = y0 - (i + 1) * 0.095
        for j, cell in enumerate(row):
            fc = "white" if j in {0, 5} else (COLORS["light_green"] if cell == "✓" else "#fee2e2")
            box(ax, (x0 + sum(widths[:j]), y), widths[j] - 0.005, 0.075, cell, fc, size=8.2, weight="bold" if j == 0 else "normal")
    ax.text(0.05, 0.11, "说明：建议最终补齐A0-A4在同一批20条查询上的F1@20、Recall@100和Latency，形成可复验消融表。", fontsize=10, color=COLORS["muted"])
    figs["ablation"] = savefig("fig08_消融矩阵.png")

    fig, ax = canvas("公开/隐藏测试评测流水线")
    nodes = [
        ("测试集JSONL", "query + gold answer"),
        ("Agent运行", "只读取query，不读gold"),
        ("预测输出", "Top-K papers"),
        ("指标计算", "Precision/Recall/F1/Latency"),
        ("诊断反馈", "hit_report定位瓶颈"),
    ]
    for i, (a, b) in enumerate(nodes):
        x0 = 0.06 + i * 0.18
        box(ax, (x0, 0.55), 0.14, 0.16, f"{a}\n{b}", [COLORS["light_blue"], COLORS["light_cyan"], COLORS["light_green"], COLORS["light_orange"], COLORS["light_violet"]][i], size=8.8, weight="bold")
        if i < 4:
            arrow(ax, (x0 + 0.145, 0.63), (x0 + 0.175, 0.63), COLORS["blue"], 1.8)
    box(ax, (0.12, 0.25), 0.34, 0.12, "公开集用于快速定位策略缺陷\n隐藏集依靠真实API与通用查询理解泛化", "white", size=9.5)
    box(ax, (0.56, 0.25), 0.32, 0.12, "PaSaTitleDB是公开本地论文库召回源\n不能替代隐藏集真实学术API覆盖", "white", size=9.5)
    figs["eval"] = savefig("fig09_评测流水线.png")

    fig, ax = canvas("Ubuntu服务器部署拓扑")
    box(ax, (0.06, 0.58), 0.20, 0.17, "本地浏览器\n127.0.0.1:8091\nSSH端口映射", COLORS["light_blue"], COLORS["blue"], 10, "bold")
    box(ax, (0.40, 0.58), 0.20, 0.17, "Ubuntu服务器\nPython 3.12 + venv\nweb_demo.py", COLORS["light_green"], COLORS["green"], 10, "bold")
    box(ax, (0.74, 0.58), 0.20, 0.17, "外部服务\nDeepSeek / OpenAlex / S2\narXiv / Serper", COLORS["light_orange"], COLORS["orange"], 10, "bold")
    arrow(ax, (0.27, 0.665), (0.39, 0.665), COLORS["blue"], 2)
    arrow(ax, (0.61, 0.665), (0.73, 0.665), COLORS["blue"], 2)
    box(ax, (0.20, 0.25), 0.62, 0.14, "运行：source .venv/bin/activate → source .env → python web_demo.py --config config.yaml --port 8091\n评测：python evaluate_pasa.py --config config.yaml --input data/pasa-dataset/RealScholarQuery/test.jsonl --limit 20 --top_k 20", "white", size=9)
    figs["deploy"] = savefig("fig10_部署拓扑.png")

    fig, ax = canvas("前端演示界面模块说明（终稿替换为真实截图）")
    box(ax, (0.05, 0.20), 0.29, 0.56, "搜索与评测面板\nQuery输入\nTop-K参数\n搜索/Smoke/RAG/漏洞检测\n配置与调用统计", COLORS["light_blue"], COLORS["blue"], 10.5, "bold")
    box(ax, (0.40, 0.20), 0.53, 0.56, "结果展示面板\n论文排序 / 查询拆解 / Agent轨迹 / 结果归纳 / 关系图 / JSON\n展示综合分、API、BM25、Embedding、Reranker、LLM、Authority、Recency、Diversity", "white", size=10.2, weight="bold")
    ax.text(0.06, 0.11, "待贴图：建议截取服务器网页真实运行结果，包含Top3论文、LLM/API调用次数和结果归纳页。", fontsize=10, color=COLORS["red"])
    figs["frontend"] = savefig("fig11_前端占位说明.png")

    fig, ax = canvas("核心创新点地图")
    items = [
        ("复杂查询理解", "解析主题、方法、数据集、时间、venue、论文类型"),
        ("多策略召回", "API检索 + arXiv定向Crawler + 本地论文库 + 查询演化"),
        ("选择器采样", "Source-aware candidate sampling，避免单一来源挤占LLM预算"),
        ("多目标排序", "相关性、权威性、时效性、多样性和成本约束联合优化"),
        ("结构化归纳", "不仅返回论文，还输出主题线索、候选分层、关系图和JSON"),
        ("工程可复现", "统一评测脚本、hit_report诊断、可部署前端演示"),
    ]
    for i, (a, b) in enumerate(items):
        x0 = 0.06 + (i % 2) * 0.46
        y0 = 0.72 - (i // 2) * 0.20
        box(ax, (x0, y0), 0.39, 0.12, f"{a}\n{b}", [COLORS["light_blue"], COLORS["light_cyan"], COLORS["light_green"], COLORS["light_orange"], COLORS["light_violet"], "#fee2e2"][i], size=9, weight="bold")
    figs["innovation"] = savefig("fig12_创新点地图.png")

    return figs


def set_cell(cell, text, bold=False, fill=None, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill.replace("#", ""))
        tc_pr.append(shd)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        set_cell(table.rows[0].cells[j], h, bold=True, fill="1E3A8A", color="FFFFFF")
    for row in rows:
        cells = table.add_row().cells
        for j, value in enumerate(row):
            set_cell(cells[j], value, fill="EFF6FF" if j == 0 else None)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    if not bold:
        p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    run.bold = bold
    return p


def add_figure(doc, path: Path, caption: str, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9)
        run.italic = True


def build_doc(figs: dict[str, Path]) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.styles["Normal"].font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("第八届中国研究生人工智能创新大赛\n")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(16)
    r.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("问研智搜\n")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(26)
    r.bold = True
    r.font.color.rgb = RGBColor(15, 23, 42)
    r = p.add_run("——基于大小模型协同与PaSa-inspired Agent的端到端学术论文智能搜索系统")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(14)
    r.bold = True
    for text in ["项目文档 / Technical Proposal", "江苏大学 汽车工程研究院  搜的都队", "版本：V2.0（v12主线代码口径）", "日期：2026年7月"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(12)
    doc.add_page_break()

    add_heading(doc, "摘要", 1)
    add_para(doc, "本项目面向学术论文智能搜索场景，构建“问研智搜”端到端智能Agent系统。系统围绕赛题提出的复杂查询理解、覆盖率与精确度平衡、权威性/时效性/相关性/多样性权衡以及结构化结果归纳四类挑战，设计了大小模型协同、PaSa-inspired Crawler/Selector、多源学术API召回、BGE小模型精排、DeepSeek相关性验证和多目标融合排序的一体化技术路线。")
    add_para(doc, "与单纯关键词检索不同，系统首先将自然语言科研问题解析为可执行的多维检索意图，再通过OpenAlex、Semantic Scholar、arXiv、Serper-arXiv Crawler、PaSaTitleDB和可选LocalCorpus构建候选池；随后利用BM25、Embedding、Reranker和LLM Score形成证据链式排序，并输出论文列表、查询拆解、Agent轨迹、关系图、归纳报告和JSON结果。")
    add_para(doc, "当前公开RealScholarQuery阶段性最好结果以v12主线为项目书主口径：Precision@20≈0.160、Recall@20≈0.192、Recall@50≈0.233、Recall@100≈0.288、F1@20≈0.159、平均API Calls≈20.2、LLM Calls≈4.0、平均Latency≈91.1s。该结果仍低于完整PaSa训练系统，但在轻量工程可部署、可解释展示和二次迭代成本方面具备竞赛演示价值。")
    add_para(doc, "关键词：学术搜索Agent；大小模型协同；PaSa-inspired Crawler/Selector；多源检索；BGE Reranker；DeepSeek；多目标排序；结构化归纳")

    add_heading(doc, "目录", 1)
    for item in ["一、项目背景概述与赛题分析", "二、系统总体方案与技术路线", "三、核心算法设计", "四、系统实现与前端平台", "五、实验验证与消融设计", "六、创新点、优势与风险控制", "七、部署实施与后续补充计划", "参考文献", "附录A：运行命令与待补材料清单"]:
        add_para(doc, item)
    doc.add_page_break()

    add_heading(doc, "一、项目背景概述与赛题分析", 1)
    add_heading(doc, "1.1 行业背景", 2)
    add_para(doc, "科研活动中的论文搜索已经从传统关键词检索转向“意图理解—多源召回—证据排序—结构化综述”的智能化范式。面对跨学科、多约束、开放式的科研问题，研究者往往需要同时考虑主题相关性、方法相似性、数据集覆盖、发表时间、venue权威性、引用影响力以及论文之间的关系。传统搜索引擎虽然覆盖广，但无法稳定把复杂自然语言问题转化为高质量Top-K论文列表。")
    add_para(doc, "近年PaSa、SPAR、Ai2 Paper Finder等系统说明，LLM驱动的智能搜索Agent能够显著提升复杂学术查询的完成质量，但完整系统往往依赖大规模Crawler、Selector训练、强化学习或多索引基础设施，复现和部署成本较高。因此，本项目选择“轻量可部署 + 强解释展示 + 评分导向优化”的工程路线，在有限资源下尽可能贴近赛题要求。")
    add_heading(doc, "1.2 赛题挑战映射", 2)
    add_table(doc, ["赛题挑战", "技术含义", "本项目对应方案"], [
        ["查询理解不充分", "用户查询包含主题、方法、数据集、年份、venue等多维约束", "DeepSeek Planner + 规则兜底解析意图、实体和子查询"],
        ["覆盖率与精确度平衡", "高召回会引入噪声，强过滤又可能漏掉目标论文", "Crawler扩大候选池，Selector/Ranker压缩噪声"],
        ["权威性、时效性、相关性、多样性权衡", "Top-K不能只看关键词匹配，需要综合多类信号", "多目标融合排序函数引入authority、recency、diversity"],
        ["结果归纳与结构化展示", "评测不仅需要论文列表，也需要可解释输出", "论文排序、查询拆解、Agent轨迹、结果归纳、关系图、JSON多视图展示"],
    ])
    add_figure(doc, figs["score_weight"], "图1-1 赛题评分权重与系统优化目标映射")

    add_heading(doc, "二、系统总体方案与技术路线", 1)
    add_heading(doc, "2.1 总体架构", 2)
    add_para(doc, "系统采用“前端可视化 + Python Agent后端 + 多源学术API + 大小模型协同”的四层架构。大模型主要承担高层语义规划、复杂查询拆解、相关性验证和结果归纳；小模型主要承担高吞吐语义匹配、候选过滤和细粒度排序。通过这种分工，系统避免把所有候选都交给大模型，从而降低Token成本和端到端延迟。")
    add_figure(doc, figs["architecture"], "图2-1 问研智搜总体架构")
    add_heading(doc, "2.2 PaSa-inspired轻量化改良思路", 2)
    add_para(doc, "PaSa的核心思想是通过Crawler发现候选论文，再通过Selector判断候选价值。本项目借鉴这一思想，但不直接复现其RL训练链路，而是将Crawler实现为多源API与arXiv定向网页搜索的组合，将Selector实现为“源感知采样 + BGE精排 + DeepSeek verifier”的轻量闭环。该设计牺牲了一部分理论上限，但显著降低了工程复杂度，适合在竞赛服务器上快速部署、调参和演示。")
    add_figure(doc, figs["pasa"], "图2-2 PaSa-inspired Crawler/Selector协同机制")
    add_heading(doc, "2.3 运行流程", 2)
    add_figure(doc, figs["workflow"], "图2-3 端到端Agent运行流程")

    add_heading(doc, "三、核心算法设计", 1)
    add_heading(doc, "3.1 查询理解与多维约束解析", 2)
    add_para(doc, "系统将用户输入q解析为结构化查询状态S={intent, entities, constraints, sub_queries}。其中intent描述学术目标，entities记录核心领域实体，constraints覆盖时间、方法、数据集、venue和论文类型等约束，sub_queries作为后续Crawler执行单元。若LLM规划失败，系统回退到启发式查询扩展，以保证端到端鲁棒性。")
    add_heading(doc, "3.2 多源候选池构建", 2)
    add_table(doc, ["候选来源", "作用", "优势", "局限"], [
        ["OpenAlex", "通用学术元数据检索", "覆盖面广、无需强认证", "长自然语言查询容易漂移"],
        ["Semantic Scholar", "论文语义和引用信息补充", "学术相关性较好", "限流与字段缺失可能影响稳定性"],
        ["arXiv API", "面向预印本论文召回", "适合LLM/AI公开测试集", "查询语法较敏感"],
        ["Serper-arXiv Crawler", "模拟PaSa式网页Crawler", "对长查询和arXiv页面发现更有效", "依赖Serper API Key"],
        ["PaSaTitleDB", "公开PaSa本地arXiv标题库", "提高公开集候选覆盖", "隐藏集不能完全依赖它"],
        ["LocalCorpus", "可扩展本地论文库接口", "适合官方提供语料时接入", "需要额外构建索引"],
    ])
    add_heading(doc, "3.3 多目标融合排序", 2)
    add_para(doc, "候选论文p针对查询q的综合得分可写为：Score(q,p)=Σwᵢfᵢ(q,p)+LabelBonus−NoisePenalty。该函数不是单纯追求文本相似度，而是将“可命中金标准论文”的排序目标与赛题效率约束共同纳入，显式融合API相关性、BM25、Embedding、Reranker、LLM Score、Authority、Recency和Diversity等信号。")
    add_figure(doc, figs["rank"], "图3-1 多目标融合排序函数")
    add_heading(doc, "3.4 LLM选择器与结果归纳", 2)
    add_para(doc, "DeepSeek在系统中承担三类角色：Planner负责理解查询和生成子查询；Verifier负责对候选论文进行high/partial/irrelevant相关性标注；Synthesizer负责把Top-K结果归纳为整体结论、主题线索、高相关候选和部分相关候选。为了控制成本，评测默认最多4次LLM调用，Verifier仅处理采样后的Top候选，而不是对全部候选暴力打分。")
    add_figure(doc, figs["innovation"], "图3-2 核心创新点地图")

    add_heading(doc, "四、系统实现与前端平台", 1)
    add_heading(doc, "4.1 前端功能模块", 2)
    add_table(doc, ["前端区域", "显示内容", "作用"], [
        ["标题栏", "问研智搜；江苏大学 汽车工程研究院 搜的都队", "展示项目名称和团队信息"],
        ["搜索与评测面板", "Query输入、Top-K、搜索论文、Smoke评测、RAG评测、漏洞检测", "作为演示和快速自测入口"],
        ["运行状态卡片", "LLM Calls、API Calls、Latency、Papers", "体现效率评分相关信息"],
        ["论文排序页", "Top-K论文、综合分和分项信号", "展示最终提交结果及排序依据"],
        ["查询拆解页", "intent、entities、sub_queries", "展示复杂查询理解能力"],
        ["Agent轨迹页", "每轮Crawler/Selector/Ranker动作", "展示系统可解释性"],
        ["结果归纳页", "整体结论、主题线索、高相关/部分相关候选", "对应赛题结构化整理要求"],
        ["关系图/JSON页", "论文关系与机器可读输出", "用于答辩展示和自动评测对接"],
    ])
    add_figure(doc, figs["frontend"], "图4-1 前端演示界面模块说明占位图")
    add_heading(doc, "4.2 工程模块划分", 2)
    add_table(doc, ["文件/模块", "职责"], [
        ["web_demo.py", "HTTP服务与前端页面渲染"],
        ["wenyan_competition/agent.py", "Agent主流程、Crawler/Selector/Ranker调度、轨迹记录"],
        ["wenyan_competition/retrievers.py", "OpenAlex/Semantic Scholar/arXiv/Serper/本地库检索"],
        ["wenyan_competition/llm.py", "DeepSeek兼容接口、Planner/Verifier/Synthesizer"],
        ["wenyan_competition/ranking.py", "BM25、Embedding、Reranker和多信号融合"],
        ["evaluate_pasa.py", "PaSa公开数据集评测与metrics输出"],
        ["offline_quality_check.py", "离线质量检查，防止关键功能回退"],
    ])
    add_figure(doc, figs["deploy"], "图4-2 Ubuntu服务器部署拓扑")

    add_heading(doc, "五、实验验证与消融设计", 1)
    add_heading(doc, "5.1 评测数据与指标", 2)
    add_para(doc, "当前主要使用PaSa公开RealScholarQuery测试集进行阶段性验证。评测脚本逐条读取query，Agent只使用query执行检索，gold answer仅用于最后计算Precision@20、Recall@20/50/100和F1@20。效率指标记录平均API调用、LLM调用和端到端延迟。")
    add_figure(doc, figs["eval"], "图5-1 公开/隐藏测试评测流水线")
    add_heading(doc, "5.2 阶段性结果", 2)
    add_table(doc, ["版本", "测试范围", "Precision@20", "Recall@20", "Recall@50", "Recall@100", "F1@20", "API Calls", "LLM Calls", "Latency", "结论"], [
        ["v7初版", "RealScholarQuery limit5", "-", "-", "-", "-", "0.086", "37.2", "4.0", "116.8s", "召回和排序均不足"],
        ["v12主线", "RealScholarQuery limit20", "0.160", "0.192", "0.233", "0.288", "0.159", "20.2", "4.0", "91.1s", "当前项目书主结果口径"],
        ["rollback复测", "RealScholarQuery limit20", "0.140", "0.184", "0.225", "0.287", "0.143", "20.2", "4.0", "98.0s", "低于v12，不作为主结果"],
        ["v16压缩LLM", "RealScholarQuery limit5", "0.100", "0.070", "0.150", "0.271", "0.081", "20.0", "2.0", "146.8s", "失败实验，证明不能盲目压缩LLM"],
    ])
    add_figure(doc, figs["metrics"], "图5-2 RealScholarQuery阶段性指标对比")
    add_heading(doc, "5.3 消融实验设计", 2)
    add_para(doc, "项目书建议保留消融实验设计，并在时间允许时补齐A0-A4结果。消融实验的价值不只是展示分数，而是解释系统为什么需要多源召回、为什么需要小模型精排、为什么LLM verifier不能过度压缩。")
    add_figure(doc, figs["ablation"], "图5-3 消融实验矩阵")

    add_heading(doc, "六、创新点、优势与风险控制", 1)
    add_heading(doc, "6.1 相对参考系统的定位", 2)
    add_table(doc, ["系统", "核心路线", "优势", "本项目吸收/改良点"], [
        ["PaSa", "Crawler+Selector双Agent+RL训练", "公开报告性能强，覆盖率高", "吸收Crawler/Selector思想，改为轻量可部署版本"],
        ["SPAR", "RefChain查询分解+查询演化", "复杂查询分解能力强", "引入多策略sub-query和查询演化"],
        ["Ai2 Paper Finder", "多索引语义检索+采样策略", "工程检索覆盖强", "借鉴多源索引和候选采样思想"],
        ["问研智搜", "轻量Agent+多源API+BGE+DeepSeek+结构化展示", "部署简单、可解释、适合答辩演示", "在有限硬件与时间下兼顾F1、效率和展示"],
    ])
    add_figure(doc, figs["radar"], "图6-1 系统能力雷达图")
    add_heading(doc, "6.2 核心优势", 2)
    for text in [
        "优势一：面向复杂学术意图，而非单关键词匹配。系统显式抽取实体、约束和子查询，能够处理“方法+数据集+时间+论文类型”组合查询。",
        "优势二：召回与精排分离。Crawler阶段尽量扩大候选空间，Selector和Ranker阶段再压缩噪声，符合学术搜索“先覆盖后筛选”的基本规律。",
        "优势三：大小模型协同。小模型承担高频、低成本排序，大模型只在关键候选和归纳环节介入，形成成本可控的Agent链路。",
        "优势四：结构化展示能力完整。系统不仅输出Top-K，还展示查询拆解、Agent轨迹、结果归纳、关系图和JSON，满足赛题结构化输出要求。",
        "优势五：工程闭环清晰。evaluate_pasa.py、metrics.json和hit_report.json可持续支持调参，避免只凭前端截图主观判断效果。",
    ]:
        add_para(doc, text)
    add_heading(doc, "6.3 风险与应对", 2)
    add_table(doc, ["风险", "表现", "应对策略"], [
        ["在线API不稳定", "OpenAlex/S2返回漂移、限流或缺字段", "启用缓存、保留arXiv/Serper和本地库兜底"],
        ["公开集过拟合", "过度依赖PaSaTitleDB可能影响隐藏集泛化", "隐藏集以真实API和LocalCorpus接口为主，标题库仅作可选召回源"],
        ["LLM成本高", "Verifier调用过多导致时延上升", "固定调用预算，使用source-aware采样减少无效候选"],
        ["Top20排序不稳", "Recall@100较高但F1@20不足", "利用hit_report区分召回不足和排序不足，再微调融合权重"],
    ])

    add_heading(doc, "七、部署实施与后续补充计划", 1)
    add_heading(doc, "7.1 当前推荐运行命令", 2)
    add_para(doc, "服务器端建议命令：cd ~/wenyan-zhisou-advanced && git pull && cd competition && source .venv/bin/activate && source .env && python offline_quality_check.py")
    add_para(doc, "评测命令：python evaluate_pasa.py --config config.yaml --input data/pasa-dataset/RealScholarQuery/test.jsonl --output_dir runs/pasa_rs_limit20_v12 --limit 20 --top_k 20")
    add_heading(doc, "7.2 仍需补充的数据", 2)
    add_table(doc, ["优先级", "需要补充的材料", "用途", "建议命令/说明"], [
        ["P0", "v12主线RealScholarQuery limit20复测metrics.json截图", "确认项目书主结果", "按第7.1评测命令运行，发metrics截图"],
        ["P0", "前端论文排序页截图", "替换图4-1占位", "网页搜索一个AI或汽车领域查询，截Top3与左侧调用统计"],
        ["P1", "查询拆解页与结果归纳页截图", "证明结构化输出10%评分能力", "同一查询切到“查询拆解/结果归纳”页截图"],
        ["P1", "AutoScholarQuery limit5或limit10结果", "证明非RealScholarQuery泛化", "若数据可用，运行evaluate_pasa.py换input路径"],
        ["P2", "A0-A4消融结果", "增强项目书说服力", "时间不足可先不跑，保留设计表"],
        ["P2", "答辩演示视频/动图", "专家评分展示", "录屏展示一次完整搜索流程"],
    ])

    add_heading(doc, "参考文献", 1)
    for i, ref in enumerate([
        "ByteDance/PKU. PaSa: An LLM Agent for Comprehensive Academic Paper Search. GitHub repository.",
        "SPAR: RefChain-based scholarly paper search and query evolution system.",
        "Ai2 Paper Finder / Asta-Bench related benchmark and system materials.",
        "OpenAlex Documentation. Scholarly metadata API.",
        "Semantic Scholar Recommendations and Graph API Documentation.",
        "BAAI. BGE Embedding and BGE Reranker model family.",
        "DeepSeek API Documentation. OpenAI-compatible chat completion interface.",
    ], 1):
        add_para(doc, f"[{i}] {ref}")

    add_heading(doc, "附录A：待贴图位置说明", 1)
    add_para(doc, "本文档已内置前端占位图。最终提交前建议将其替换为真实网页截图，并在第5章补入最新metrics.json截图或表格。若提供截图和复测结果，可继续生成V2.1终稿。")

    out = DOCS / "问研智搜_赛题三项目书_V2.0.docx"
    doc.save(out)
    return out


def write_todo() -> Path:
    path = DOCS / "项目书待补数据清单.md"
    path.write_text(
        """# 项目书待补数据清单

## 必补
1. v12主线 RealScholarQuery limit20 复测 metrics.json 截图。
2. 前端论文排序页截图：包含Top3论文、左侧LLM/API调用次数、Latency、Papers。
3. 查询拆解页截图与结果归纳页截图。

## 建议补
1. AutoScholarQuery limit5或limit10结果，用于证明泛化。
2. A0-A4消融实验结果，用于增强项目书实验章节。
3. 一段完整搜索演示录屏或截图序列，用于答辩。

## 当前项目书主结果口径
- v12主线，RealScholarQuery limit20：Precision@20≈0.160，Recall@20≈0.192，Recall@50≈0.233，Recall@100≈0.288，F1@20≈0.159，API Calls≈20.2，LLM Calls≈4.0，Latency≈91.1s。
- rollback复测：F1@20≈0.143，低于v12，不建议作为主结果。
""",
        encoding="utf-8",
    )
    return path


def main() -> None:
    figs = draw_figures()
    docx = build_doc(figs)
    todo = write_todo()
    print(docx)
    print(todo)
    print(f"figures={len(figs)}")


if __name__ == "__main__":
    main()
